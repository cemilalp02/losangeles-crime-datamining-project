from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "reports" / "Crime_Forecast_LA_Final_Report.docx"
FIG = ROOT / "reports" / "figures"
DOCS_SKILL = Path(
    r"C:\Users\cemil\.codex\plugins\cache\openai-primary-runtime\documents\26.430.10722\skills\documents\scripts"
)
sys.path.insert(0, str(DOCS_SKILL))
from table_geometry import apply_table_geometry, column_widths_from_weights  # noqa: E402


ACCENT = RGBColor(31, 78, 95)
MUTED = RGBColor(89, 89, 89)
LIGHT_SHADE = "F3F6F8"
HEADER_SHADE = "E7EEF2"
CONTENT_WIDTH_DXA = 9360


def set_run_font(run, name="Arial", size=None, bold=None, italic=None, color=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_bottom_border(paragraph, color="B8C4CC", size="6"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_page_field(paragraph, field_name: str):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {field_name} "
    run._r.append(instr)

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_text(cell, text, bold=False, size=9.2, color=None, align=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(str(text))
    set_run_font(run, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def format_table(table, widths, header=True, font_size=9.2, center_cols=None):
    table.style = "Table Grid"
    table.autofit = False
    apply_table_geometry(table, widths, table_width_dxa=sum(widths), indent_dxa=0)
    center_cols = set(center_cols or [])
    for row_idx, row in enumerate(table.rows):
        if row_idx == 0 and header:
            repeat_table_header(row)
        for col_idx, cell in enumerate(row.cells):
            if row_idx == 0 and header:
                set_cell_shading(cell, HEADER_SHADE)
                for p in cell.paragraphs:
                    for r in p.runs:
                        set_run_font(r, size=font_size, bold=True, color=ACCENT)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                for p in cell.paragraphs:
                    p.paragraph_format.space_after = Pt(0)
                    if col_idx in center_cols:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for r in p.runs:
                        set_run_font(r, size=font_size)


def add_caption(doc, text):
    p = doc.add_paragraph(style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, size=9.3, italic=True, color=MUTED)
    return p


def add_note_box(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    widths = column_widths_from_weights([1], CONTENT_WIDTH_DXA)
    apply_table_geometry(table, widths, table_width_dxa=CONTENT_WIDTH_DXA, indent_dxa=0)
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_SHADE)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    set_run_font(r, size=10.5, bold=True, color=ACCENT)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    set_run_font(r2, size=9.8)


def add_table(doc, rows, weights, caption=None, font_size=9.2, center_cols=None):
    if caption:
        add_caption(doc, caption)
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    for i, row in enumerate(rows):
        for j, text in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.CENTER if j in set(center_cols or []) else None
            set_cell_text(table.cell(i, j), text, bold=(i == 0), size=font_size, align=align)
    widths = column_widths_from_weights(weights, CONTENT_WIDTH_DXA)
    format_table(table, widths, header=True, font_size=font_size, center_cols=center_cols)
    doc.add_paragraph()
    return table


def add_figure(doc, image_name, caption, width_inches=6.2):
    path = FIG / image_name
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width_inches))
    add_caption(doc, caption)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(item)
        set_run_font(r, size=11)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(item)
        set_run_font(r, size=11)


def add_para(doc, text, keep_with_next=False):
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = keep_with_next
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_run_font(r, size=11)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = True
    return p


def setup_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.08
    normal.paragraph_format.space_after = Pt(6)

    title = styles["Title"]
    title.font.name = "Arial"
    title._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    title._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    title.font.size = Pt(24)
    title.font.bold = True
    title.font.color.rgb = ACCENT
    title.paragraph_format.space_after = Pt(8)

    subtitle = styles["Subtitle"]
    subtitle.font.name = "Arial"
    subtitle._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    subtitle._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    subtitle.font.size = Pt(12)
    subtitle.font.color.rgb = MUTED

    for name, size in [("Heading 1", 16), ("Heading 2", 13.5), ("Heading 3", 11.5)]:
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = ACCENT
        style.paragraph_format.space_before = Pt(12 if name == "Heading 1" else 8)
        style.paragraph_format.space_after = Pt(5)

    caption = styles["Caption"]
    caption.font.name = "Arial"
    caption._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = MUTED

    header = section.header
    hp = header.paragraphs[0]
    hp.text = ""
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hr = hp.add_run("Crime Forecast LA | Data Mining Final Report")
    set_run_font(hr, size=8.5, color=MUTED)
    add_bottom_border(hp)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fr = fp.add_run("Crime Forecast LA | Final Report")
    set_run_font(fr, size=8.5, color=MUTED)


def build_report():
    doc = Document()
    setup_document(doc)
    cp = doc.core_properties
    cp.title = "Crime Forecast LA - Final Report"
    cp.subject = "Spatio-temporal crime forecasting, hotspot discovery, and explainable risk modeling"
    cp.author = "Crime Forecast LA Team"

    # 1. Cover Page
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Crime Forecast LA")

    p = doc.add_paragraph(style="Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Spatio-Temporal Crime Forecasting, Hotspot Discovery and Explainable Risk Modeling")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Final Report for Data Mining (DM) Course")
    set_run_font(r, size=13, bold=True, color=MUTED)

    doc.add_paragraph()
    add_note_box(
        doc,
        "Project Scope",
        "This report studies LAPD Crime Data from 2020 to Present with an end-to-end data mining pipeline for classification, clustering, forecasting, risk scoring, and model explainability.",
    )
    doc.add_paragraph()

    cover_rows = [
        ["Required Cover Item", "Submitted Information"],
        ["Title of the Project", "Crime Forecast LA: Spatio-Temporal Crime Forecasting, Hotspot Discovery and Explainable Risk Modeling"],
        ["Course Name", "Data Mining (DM) Course"],
        ["Team Information", "Omitted from the public repository"],
        ["Date of Submission", "16 May 2026"],
        ["Dataset", "LAPD Crime Data from 2020 to Present, accessed through Kaggle and the City of Los Angeles open-data ecosystem"],
    ]
    add_table(doc, cover_rows, [1.4, 3.6], font_size=9.5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Repository reviewed: dataminingproject-01")
    set_run_font(r, size=10, italic=True, color=MUTED)
    doc.add_page_break()

    # 2. Abstract
    add_heading(doc, "2. Abstract", 1)
    abstract = (
        "Crime incidents in large metropolitan areas form a high-volume, spatially and temporally dependent data mining problem. "
        "This project analyzes approximately 955,339 Los Angeles Police Department crime records from 2020 to 2024 to identify where crime risk is concentrated, how incident characteristics relate to violent crime and arrest outcomes, and how short-term area-level crime volume can be forecast. "
        "The experimental pipeline performs schema-aware loading, data cleaning, exploratory analysis, incident-level feature engineering, supervised classification, density-based hotspot discovery, global LightGBM forecasting, composite risk scoring, and SHAP-based model interpretation. "
        "For classification, Logistic Regression and LightGBM are compared on a stratified 80/20 split of a 200,000-row sample. "
        "For forecasting, LightGBM regressors are evaluated with forward-chained validation and compared against rolling-mean naive baselines. "
        "The best violent-crime classifier reaches 0.923 ROC-AUC and 0.620 PR-AUC, while the arrest model is more difficult because the dataset lacks case-processing context. "
        "The daily and weekly forecasters improve MAE over naive baselines by 9.8 percent and 15.2 percent, respectively. "
        "The final risk score ranks LAPD areas by recent intensity, trend, violence share, and forecasted volume. "
        "The resulting system is suitable for analytical decision support, hotspot monitoring, and transparent public-safety resource planning."
    )
    wc = len(re.findall(r"\b[\w'-]+\b", abstract))
    assert 150 <= wc <= 250, f"Abstract must be 150-250 words, got {wc}"
    add_para(doc, abstract)

    # 3. Introduction
    add_heading(doc, "3. Introduction", 1)
    add_heading(doc, "3.1 Problem Description", 2)
    add_para(
        doc,
        "The project addresses a practical data mining question: given hundreds of thousands of crime reports, can we discover stable temporal, spatial, and categorical patterns; classify important incident outcomes; and forecast near-term risk at the area level? The dataset contains incident timestamps, LAPD area identifiers, crime descriptions, victim attributes, weapon indicators, status codes, and coordinates, making it appropriate for supervised learning, clustering, forecasting, and explainability tasks.",
    )
    add_heading(doc, "3.2 Motivation", 2)
    add_para(
        doc,
        "Crime analysis is challenging because incidents are sparse at fine spatial scales, delayed in reporting, class-imbalanced, and affected by non-observed social and investigative variables. A data mining solution must therefore combine prediction with interpretable pattern discovery rather than report a single accuracy number. The motivation is to turn raw incident logs into reproducible evidence that supports area-level monitoring, tactical planning, and critical evaluation of model limits.",
    )
    add_heading(doc, "3.3 Approach and Report Organization", 2)
    add_para(
        doc,
        "The project implements a modular Python pipeline. Raw data are cleaned and transformed into incident-level and area-time feature stores. The incident store supports violent-crime and arrest classification; the area-time panels support hotspot discovery, forecasting, and risk scoring. This report follows the required final-report structure: related work, algorithms and methodology, experimental setup, evaluation, conclusions, and references.",
    )

    # 4. Background and Related Work
    add_heading(doc, "4. Background and Related Work", 1)
    add_para(
        doc,
        "Urban crime mining commonly combines four families of methods: descriptive exploratory analysis, supervised classification, spatial clustering, and time-series forecasting. Public police datasets such as the LAPD open-data release provide incident-level detail but also introduce known quality issues, including reporting delays, missing coordinates, and changes in enforcement or reporting behavior over time [1], [2].",
    )
    add_para(
        doc,
        "For spatial pattern discovery, density-based clustering is a strong fit because hotspots are irregularly shaped and can include noise points. DBSCAN is especially relevant because it identifies dense spatial regions without requiring the number of clusters in advance [3]. For predictive modeling on mixed numeric and categorical tabular features, gradient-boosted decision trees remain a competitive state-of-the-art choice. LightGBM is used in this project because it scales well to large tabular datasets and supports efficient learning with many features [4].",
    )
    add_para(
        doc,
        "Interpretability is essential for crime-related applications. SHAP values provide a consistent feature-attribution framework that helps explain model outputs while keeping the distinction between statistical association and causal explanation clear [5]. The implementation also relies on established scientific Python tools, especially scikit-learn for preprocessing and baseline models and pandas for data manipulation [6], [7].",
    )

    # 5. Algorithms and Methodology
    add_heading(doc, "5. Algorithms and Methodology", 1)
    add_heading(doc, "5.1 Data Mining Tasks and Algorithms", 2)
    method_rows = [
        ["Task", "Algorithm or Technique", "Reason for Selection", "Primary Output"],
        ["Exploratory Data Analysis", "Time-series aggregation, categorical frequency analysis, spatial summaries", "Reveals temporal seasonality, area concentration, crime mix, class imbalance, and data quality issues before modeling.", "Static figures and dataset overview JSON"],
        ["Violent Crime Classification", "Logistic Regression and LightGBM classifier", "Logistic Regression gives an interpretable linear baseline; LightGBM captures nonlinear interactions among time, location, victim, and weapon features.", "Accuracy, F1, ROC-AUC, PR-AUC, SHAP plot"],
        ["Arrest Outcome Classification", "Logistic Regression and LightGBM classifier", "Arrest is a difficult imbalanced target; comparing a baseline and boosted trees shows whether nonlinear tabular learning adds value.", "Accuracy, F1, ROC-AUC, PR-AUC, SHAP plot"],
        ["Hotspot Discovery", "Grid ranking, DBSCAN with haversine distance, Folium heatmaps", "Combines simple cell-level intensity with density-based spatial clusters and interactive maps.", "Top grid cells, DBSCAN clusters, HTML maps"],
        ["Crime Volume Forecasting", "Global LightGBM regression with lag and rolling-window features", "A global model shares information across 21 LAPD areas and handles nonlinear calendar and lag effects.", "Daily and weekly forecast metrics and predictions"],
        ["Area Risk Scoring", "Weighted normalized composite index", "Produces a transparent operational ranking from intensity, trend, violent share, and forecast components.", "Area risk table and tiers"],
    ]
    add_table(doc, method_rows, [1.0, 1.25, 2.1, 1.1], "Table 1. Data mining tasks, algorithms, and project outputs.", font_size=8.2)

    add_heading(doc, "5.2 Data Preparation Workflow", 2)
    add_bullets(
        doc,
        [
            "Schema-aware loading: raw LAPD fields are renamed to consistent snake_case names and converted to memory-efficient types.",
            "Temporal normalization: occurrence and report dates are parsed; hour, day of week, month, year, weekend flag, report lag, and part of day are derived.",
            "Spatial cleaning: coordinates outside a Los Angeles bounding box and zero-coordinate placeholders are set to missing and later imputed where needed for modeling.",
            "Victim and status cleaning: invalid victim ages are treated as missing; victim sex is normalized; LAPD status codes are mapped to arrest/non-arrest outcomes.",
            "Target derivation: is_violent is defined from a curated violent-crime code set, while is_arrest is derived from adult and juvenile arrest status codes.",
            "Feature-store creation: incident-level features support classification, while daily and weekly area panels support forecasting and risk scoring.",
        ],
    )

    add_heading(doc, "5.3 Feature Extraction and Knowledge Discovery Process", 2)
    add_para(
        doc,
        "Incident-level models use numeric variables such as hour, day of week, month, victim age, report lag, weapon flag, latitude, longitude, and cyclical time encodings. Categorical variables include area name, victim sex, victim descent, and part of day. Forecasting models use area-level crime counts with lag features, rolling means and standard deviations, day-of-week and month indicators, week-of-year, day-of-year, weekend flags, and area identifiers.",
    )
    add_para(
        doc,
        "The knowledge discovery process moves from descriptive evidence to predictive and prescriptive artifacts: first, EDA identifies temporal and spatial concentrations; second, classifiers quantify whether incident attributes predict violent crime and arrest outcomes; third, DBSCAN and grid ranking expose spatial hotspots; fourth, LightGBM forecasts near-term volume; and finally, a composite score transforms model outputs into a ranked area-risk view.",
    )

    add_heading(doc, "5.4 Performance Metrics", 2)
    add_para(
        doc,
        "Classification performance is evaluated with accuracy, positive-class F1, ROC-AUC, and PR-AUC. PR-AUC is especially important because both targets are imbalanced: violent crime is approximately 16.7 percent of incidents and arrest is approximately 9.1 percent. Forecasting performance is evaluated with MAE, RMSE, MAPE, R-squared, naive MAE, and skill score, where skill is defined as 1 - model MAE / naive MAE. Hotspot and risk outputs are evaluated through interpretability, spatial coherence, and consistency with EDA patterns.",
    )

    # 6. Experimental Setup
    add_heading(doc, "6. Experimental Setup", 1)
    add_heading(doc, "6.1 Dataset Description", 2)
    dataset_rows = [
        ["Property", "Value"],
        ["Source", "LAPD Crime Data from 2020 to Present, distributed through Kaggle and linked to LA open data"],
        ["Number of raw samples", "955,339 crime incidents"],
        ["Raw features", "28 columns"],
        ["Features after preprocessing", "41 columns in the cleaned analytical table"],
        ["Date range used", "2020-01-01 to 2024-06-24"],
        ["Spatial coverage", "21 LAPD areas"],
        ["Distinct crime types", "139"],
        ["Target variables", "is_violent and is_arrest"],
        ["Class imbalance", "Violent share: 16.73 percent; arrest share: 9.06 percent"],
        ["Missing or invalid coordinates", "Approximately 0.24 percent after LA bounding-box checks"],
    ]
    add_table(doc, dataset_rows, [1.3, 3.7], "Table 2. Dataset profile used in the experiments.", font_size=9.0)

    add_heading(doc, "6.2 Pre-processing Steps", 2)
    add_bullets(
        doc,
        [
            "Rows with missing occurrence dates are removed because temporal order is central to the analysis.",
            "Victim age values less than or equal to zero or greater than 110 are replaced with missing values and imputed for modeling.",
            "Invalid latitude and longitude values are masked using the configured Los Angeles bounding box.",
            "Categorical variables are encoded with one-hot encoding for scikit-learn pipelines; rare categories are controlled with min_frequency = 200.",
            "Numeric variables are standardized for Logistic Regression and processed through a ColumnTransformer shared across classification experiments.",
            "Forecasting data are trimmed for incomplete reporting tails before model training and risk scoring.",
        ],
    )

    add_heading(doc, "6.3 Experimental Design", 2)
    setup_rows = [
        ["Component", "Design Choice"],
        ["Classification split", "Stratified 80/20 train-test split on a 200,000-row sample, random_state = 42"],
        ["Classification models", "Logistic Regression with class_weight = balanced; LightGBM classifier with 400 estimators, learning rate 0.05, num_leaves 63"],
        ["Forecast validation", "Forward-chained design: daily model uses 60 validation days and 60 test days; weekly model uses 8 validation weeks and 8 test weeks"],
        ["Forecast baseline", "Daily: previous 28-day area rolling mean; weekly: previous 4-week rolling mean"],
        ["Forecast model", "Global LightGBM regressor with lag, rolling-window, calendar, and area features"],
        ["Hotspot setup", "DBSCAN on sampled coordinates using haversine distance, eps = 0.4 km, min_samples = 80; grid cells are approximately 250 m by 250 m"],
        ["Risk score weights", "0.35 intensity, 0.20 trend, 0.20 violence, 0.25 forecast"],
    ]
    add_table(doc, setup_rows, [1.25, 3.75], "Table 3. Experimental setup by modeling component.", font_size=8.7)

    add_heading(doc, "6.4 Software and Hardware Environment", 2)
    add_para(
        doc,
        "The project is implemented in Python and is reproducible through pipeline.py. The README specifies Windows 10 and Python 3.10 as the tested environment. The main libraries are pandas, NumPy, scikit-learn, LightGBM, SHAP, matplotlib, seaborn, Folium, Plotly, Streamlit, and PyArrow. The modeling pipeline is CPU-based and does not require a GPU. All stochastic components use RANDOM_STATE = 42.",
    )

    # 7. Experimental Evaluation
    add_heading(doc, "7. Experimental Evaluation", 1)
    add_heading(doc, "7.1 Exploratory Results and Discovered Patterns", 2)
    add_para(
        doc,
        "The exploratory analysis confirms that crime reports have strong temporal and spatial structure. Daily crime volume changes over the observation period, late afternoon through night hours have higher counts, and incidents are concentrated in a relatively small set of LAPD areas. These patterns justify the use of time-aware features, area-level panels, and spatial hotspot discovery.",
    )
    add_figure(doc, "01_crimes_over_time.png", "Figure 1. Crime volume over time after preprocessing.", width_inches=6.2)
    add_figure(doc, "02_hour_dow_heatmap.png", "Figure 2. Hour by day-of-week heatmap showing temporal concentration.", width_inches=5.9)
    add_figure(doc, "03_top_areas.png", "Figure 3. Highest-volume LAPD areas in the dataset.", width_inches=5.9)
    add_figure(doc, "07_violent_share_by_area.png", "Figure 4. Area-level variation in violent-crime share.", width_inches=5.9)

    doc.add_page_break()
    add_heading(doc, "7.2 Classification Results", 2)
    class_rows = [
        ["Target", "Model", "Accuracy", "F1", "ROC-AUC", "PR-AUC"],
        ["is_violent", "Logistic Regression", "0.823", "0.650", "0.915", "0.593"],
        ["is_violent", "LightGBM", "0.826", "0.654", "0.923", "0.620"],
        ["is_arrest", "Logistic Regression", "0.676", "0.266", "0.712", "0.184"],
        ["is_arrest", "LightGBM", "0.683", "0.279", "0.741", "0.205"],
    ]
    add_table(doc, class_rows, [1.0, 1.35, 0.75, 0.65, 0.75, 0.75], "Table 4. Incident-level classification results.", font_size=8.8, center_cols=[2, 3, 4, 5])
    add_para(
        doc,
        "LightGBM is the best model for both targets. The improvement is modest in accuracy but clearer in ROC-AUC and PR-AUC, which are more informative under class imbalance. The violent-crime classifier is strong because crime code, weapon use, time, and area variables carry substantial signal. The arrest classifier is weaker because arrest depends on factors not captured in the incident table, including suspect availability, witness information, evidence quality, and investigative follow-up.",
    )
    add_figure(doc, "shap_is_violent.png", "Figure 5. SHAP summary for the LightGBM violent-crime classifier.", width_inches=6.0)

    add_heading(doc, "7.3 Hotspot Discovery Results", 2)
    add_para(
        doc,
        "The project uses complementary hotspot views. Grid ranking identifies dense 250 m cells and is easy to interpret operationally. DBSCAN discovers dense spatial clusters without predefining the number of hotspots and labels isolated points as noise. Folium heatmaps provide interactive inspection for all crime and violent-crime subsets. The largest discovered concentrations align with central and South Los Angeles patterns observed in the EDA figures.",
    )

    add_heading(doc, "7.4 Forecasting Results", 2)
    forecast_rows = [
        ["Granularity", "MAE", "RMSE", "R-squared", "Naive MAE", "Skill"],
        ["Daily per area", "7.79", "9.48", "-0.13", "8.64", "+9.8 percent"],
        ["Weekly per area", "40.45", "53.82", "-0.10", "47.73", "+15.2 percent"],
    ]
    add_table(doc, forecast_rows, [1.35, 0.75, 0.75, 0.85, 0.85, 1.0], "Table 5. LightGBM spatio-temporal forecasting results.", font_size=8.8, center_cols=[1, 2, 3, 4, 5])
    add_para(
        doc,
        "Both LightGBM forecasters beat their rolling-mean baselines, with the weekly model showing the stronger skill score. The negative R-squared values should not be read as total model failure; the test period has distribution shift and relatively low variance, so baseline-relative MAE is the more stable criterion. The forecasts are best interpreted as short-term decision-support estimates rather than exact counts.",
    )

    add_heading(doc, "7.5 Composite Risk Evaluation", 2)
    risk_rows = [
        ["Rank", "Area", "Risk Score", "Tier", "Intensity", "Trend z", "Violent Share", "Forecast"],
        ["1", "77th Street", "0.92", "Critical", "30.1", "-0.34", "23 percent", "33.8"],
        ["2", "Southwest", "0.82", "Critical", "29.2", "-0.28", "19 percent", "30.8"],
        ["3", "N Hollywood", "0.74", "Critical", "29.7", "-0.30", "12 percent", "29.0"],
        ["4", "Southeast", "0.69", "Critical", "23.1", "-0.35", "25 percent", "27.5"],
        ["5", "Van Nuys", "0.62", "Critical", "24.5", "-0.09", "12 percent", "25.1"],
    ]
    add_table(doc, risk_rows, [0.5, 1.2, 0.9, 0.8, 0.9, 0.75, 0.95, 0.8], "Table 6. Top five LAPD areas by composite risk score.", font_size=8.3, center_cols=[0, 2, 3, 4, 5, 6, 7])
    add_para(
        doc,
        "The top-risk areas combine high recent intensity, meaningful violent-crime share, and high forecasted volume. The negative trend values in the table reflect broader city-wide decline in the final portion of the dataset, but these areas remain highest in absolute operational risk. Because the score is a weighted index rather than a black-box prediction, its assumptions can be audited and adjusted.",
    )

    add_heading(doc, "7.6 Critical Analysis, Error Analysis, and Limitations", 2)
    add_bullets(
        doc,
        [
            "Performance: violent-crime classification is reliable enough for analytic triage, while arrest prediction has limited recall and should not be treated as an operational decision model.",
            "Interpretability: SHAP clarifies influential features but does not establish causality. Model explanations should be presented as associations learned from historical data.",
            "Data limitations: the dataset lacks suspect availability, investigative evidence, staffing, local events, weather, land use, and socioeconomic variables that could improve forecasting and arrest modeling.",
            "Temporal limitations: the final weeks of the source data are affected by reporting lag, so the pipeline trims incomplete tails before forecasting and risk scoring.",
            "Spatial limitations: LAPD area-level forecasts are stable but coarse; finer grid forecasts would be more actionable and also more sensitive to noise.",
            "Ethical limitations: crime-risk systems can reinforce historical policing patterns. A production version would require fairness review, governance rules, and human oversight.",
        ],
    )

    # 8. Conclusions and Future Work
    add_heading(doc, "8. Conclusions and Future Work", 1)
    add_heading(doc, "8.1 Summary of Work and Main Outcomes", 2)
    add_para(
        doc,
        "The project successfully converts raw LAPD crime reports into a reproducible data mining system. It creates cleaned analytical stores, explores temporal and spatial patterns, trains classification models, discovers hotspots, forecasts daily and weekly area-level crime volume, ranks areas with a transparent risk score, and explains model behavior with SHAP. The main quantitative outcomes are a 0.923 ROC-AUC violent-crime classifier, a positive-skill daily and weekly forecaster, and a ranked risk table led by 77th Street, Southwest, N Hollywood, Southeast, and Van Nuys.",
    )

    add_heading(doc, "8.2 Strengths and Weaknesses", 2)
    add_bullets(
        doc,
        [
            "Strength: the pipeline is end-to-end, modular, deterministic, and supported by both notebooks and a Streamlit dashboard.",
            "Strength: multiple data mining tasks are connected into one coherent analytical workflow instead of being treated as isolated experiments.",
            "Strength: model results are benchmarked against baselines and interpreted with SHAP, which improves technical credibility.",
            "Weakness: arrest outcome modeling is constrained by missing investigation-level variables and severe class imbalance.",
            "Weakness: forecasts are area-level and therefore not as granular as a patrol-beat or grid-cell operational product.",
        ],
    )

    add_heading(doc, "8.3 Lessons Learned", 2)
    add_para(
        doc,
        "The most important lesson is that data preparation decisions dominate downstream credibility. Reporting lag, invalid coordinates, target definition, and class imbalance directly shape the model results. A second lesson is that baseline-relative evaluation is essential for forecasting: a model can look weak under R-squared during distribution shift while still improving operational MAE over a strong naive benchmark. Finally, interpretability should be designed into the pipeline from the beginning rather than added after modeling is complete.",
    )

    add_heading(doc, "8.4 Future Work", 2)
    add_bullets(
        doc,
        [
            "Add external covariates such as weather, public events, land-use indicators, holidays, and neighborhood-level context.",
            "Move from LAPD area forecasts to 1 km grid or patrol-beat forecasts with spatial smoothing and uncertainty intervals.",
            "Evaluate additional temporal models such as Prophet, SARIMAX, temporal convolution, graph neural networks, or transformer-based sequence models.",
            "Perform fairness and bias audits before any operational use, especially for features related to victim demographics or historically policed neighborhoods.",
            "Automate weekly refreshes, model monitoring, and dashboard publication with data versioning and drift checks.",
        ],
    )

    # 9. References
    add_heading(doc, "9. References", 1)
    refs = [
        "[1] H. Alam, \"Crime Dataset,\" Kaggle. [Online]. Available: https://www.kaggle.com/datasets/haseefalam/crime-dataset. Accessed: 16 May 2026.",
        "[2] Los Angeles Police Department, \"Crime Data from 2020 to Present,\" City of Los Angeles Open Data. [Online]. Available: https://data.lacity.org/. Accessed: 16 May 2026.",
        "[3] M. Ester, H.-P. Kriegel, J. Sander, and X. Xu, \"A density-based algorithm for discovering clusters in large spatial databases with noise,\" in Proceedings of the Second International Conference on Knowledge Discovery and Data Mining, 1996, pp. 226-231.",
        "[4] G. Ke et al., \"LightGBM: A highly efficient gradient boosting decision tree,\" Advances in Neural Information Processing Systems, vol. 30, 2017.",
        "[5] S. M. Lundberg and S.-I. Lee, \"A unified approach to interpreting model predictions,\" Advances in Neural Information Processing Systems, vol. 30, 2017.",
        "[6] F. Pedregosa et al., \"Scikit-learn: Machine learning in Python,\" Journal of Machine Learning Research, vol. 12, pp. 2825-2830, 2011.",
        "[7] W. McKinney, \"Data structures for statistical computing in Python,\" in Proceedings of the 9th Python in Science Conference, 2010, pp. 56-61.",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(ref)
        set_run_font(r, size=10)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)


if __name__ == "__main__":
    build_report()
    print(OUT)
